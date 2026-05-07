"""
Создаёт reference.docx для pandoc с настройками ГОСТ 7.32-2017.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_rule=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
    if line_rule:
        spacing.set(qn('w:lineRule'), line_rule)
    # Remove existing spacing
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    pPr.append(spacing)

def set_run_font(run, name, size_pt, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

def make_reference_docx(output_path='reference.docx'):
    doc = Document()

    # --- Page setup: ГОСТ margins ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)    # ГОСТ: левое 30 мм
    section.right_margin = Cm(1.5)   # ГОСТ: правое 15 мм
    section.top_margin = Cm(2.0)     # ГОСТ: верхнее 20 мм
    section.bottom_margin = Cm(2.0)  # ГОСТ: нижнее 20 мм

    # --- Default (Normal) style ---
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(14)
    nf = normal._element.find('.//' + qn('w:rFonts'))
    # Set paragraph spacing and indent for Normal
    pPr = normal.paragraph_format._element if hasattr(normal.paragraph_format, '_element') else None
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Set font in XML for all character properties
    rPr = normal.element.find('.//' + qn('w:rPr'))

    # --- Heading 1 ---
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Heading 2 ---
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Cm(1.25)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Heading 3 ---
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(14)
    h3.font.bold = False
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = Cm(1.25)
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Caption style ---
    try:
        caption_style = doc.styles['Caption']
    except:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = 'Times New Roman'
    caption_style.font.size = Pt(12)
    caption_style.font.italic = True
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.first_line_indent = Cm(0)
    caption_style.paragraph_format.space_before = Pt(3)
    caption_style.paragraph_format.space_after = Pt(6)

    # --- Table style ---
    try:
        table_style = doc.styles['Table Grid']
        table_style.font.name = 'Times New Roman'
        table_style.font.size = Pt(12)
    except:
        pass

    # --- Add page number in footer ---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.save(output_path)
    print(f' reference.docx создан: {output_path}')

if __name__ == '__main__':
    make_reference_docx()